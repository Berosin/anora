import io

import pytest
from docx import Document

from app.document_processing.extract_docx import DocxExtractionError, extract_sections


def make_docx_bytes() -> bytes:
    doc = Document()
    doc.add_heading("Eligibility", level=1)
    doc.add_paragraph("Students must maintain a minimum aggregate of 60%.")
    doc.add_heading("Important Dates", level=1)
    doc.add_paragraph("Applications open on March 1st.")
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def test_extract_sections_tracks_nearest_heading():
    sections = extract_sections(make_docx_bytes())
    texts_by_section = {}
    for entry in sections:
        texts_by_section.setdefault(entry.section, []).append(entry.text)

    assert "Eligibility" in texts_by_section
    assert any("60%" in t for t in texts_by_section["Eligibility"])
    assert "Important Dates" in texts_by_section
    assert any("March 1st" in t for t in texts_by_section["Important Dates"])


def test_extract_sections_rejects_corrupt_file():
    with pytest.raises(DocxExtractionError):
        extract_sections(b"not a real docx")
