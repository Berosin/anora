import io

import fitz
import pytest
from docx import Document

from app.document_processing.pipeline import DocumentProcessingError, process_document


def make_pdf_bytes(pages_text: list[str]) -> bytes:
    doc = fitz.open()
    for text in pages_text:
        page = doc.new_page()
        page.insert_text((72, 72), text)
    data = doc.tobytes()
    doc.close()
    return data


def make_docx_bytes() -> bytes:
    doc = Document()
    doc.add_heading("Eligibility", level=1)
    doc.add_paragraph(
        "Students must maintain a minimum aggregate of 60 percent across all "
        "semesters to remain eligible for on-campus placements this year."
    )
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def test_pipeline_pdf_produces_page_tagged_chunks():
    pdf_bytes = make_pdf_bytes(
        [
            "Placement eligibility requires a minimum aggregate of 60 percent.",
            "Internship applications open every March and close in April.",
        ]
    )
    chunks, page_count = process_document(
        file_bytes=pdf_bytes,
        file_type="pdf",
        document_id="doc-1",
        document_name="Placement_Policy.pdf",
        knowledge_base_id="kb-1",
    )
    assert page_count == 2
    assert len(chunks) >= 2
    assert {c.metadata.page for c in chunks} == {1, 2}
    for c in chunks:
        assert c.metadata.document_id == "doc-1"
        assert c.metadata.knowledge_base_id == "kb-1"
        assert c.metadata.chunk_id  # non-empty UUID


def test_pipeline_docx_produces_section_tagged_chunks():
    chunks, page_count = process_document(
        file_bytes=make_docx_bytes(),
        file_type="docx",
        document_id="doc-2",
        document_name="Handbook.docx",
        knowledge_base_id="kb-1",
    )
    assert page_count is None
    assert len(chunks) >= 1
    assert chunks[0].metadata.section == "Eligibility"


def test_pipeline_txt_produces_chunks():
    text = ("The research grant covers travel and equipment costs. " * 30).encode("utf-8")
    chunks, page_count = process_document(
        file_bytes=text,
        file_type="txt",
        document_id="doc-3",
        document_name="Grant_Info.txt",
        knowledge_base_id="kb-1",
    )
    assert page_count is None
    assert len(chunks) >= 1


def test_pipeline_rejects_unsupported_type():
    with pytest.raises(DocumentProcessingError):
        process_document(
            file_bytes=b"whatever",
            file_type="exe",
            document_id="doc-4",
            document_name="malware.exe",
            knowledge_base_id="kb-1",
        )


def test_pipeline_rejects_corrupt_pdf():
    with pytest.raises(DocumentProcessingError):
        process_document(
            file_bytes=b"not a pdf",
            file_type="pdf",
            document_id="doc-5",
            document_name="broken.pdf",
            knowledge_base_id="kb-1",
        )
